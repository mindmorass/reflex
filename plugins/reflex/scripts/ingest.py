#!/usr/bin/env python3
"""
Ingest local files into Qdrant vector database.

Supported formats:
- PDF (.pdf)
- Markdown (.md)
- Text (.txt, .rst)
- HTML (.html, .htm)
- EPUB (.epub)
- Word (.docx)
- Jupyter notebooks (.ipynb)
- Code files (.py, .js, .ts, .go, .rs, .java, .c, .cpp, .rb, .sh)

Usage:
    uvx --with pymupdf,fastembed,qdrant-client,python-docx,ebooklib,beautifulsoup4 \
        python ingest.py <path> [--collection NAME] [--chunk-size WORDS]

Examples:
    python ingest.py ~/Documents/manual.pdf
    python ingest.py ~/notes/ --collection research
    python ingest.py ~/code/project --chunk-size 300
"""

import argparse
import hashlib
import json
import os
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Qdrant client
try:
    from qdrant_client import QdrantClient
    from qdrant_client.models import PointStruct, VectorParams, Distance
except ImportError:
    print("Install: pip install qdrant-client")
    sys.exit(1)

# Embedding model
try:
    from fastembed import TextEmbedding
except ImportError:
    print("Install: pip install fastembed")
    sys.exit(1)


# =============================================================================
# Extractors for different file formats
# =============================================================================

def extract_pdf(path: Path) -> Tuple[str, Dict]:
    """Extract text from PDF using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        print("Install: pip install pymupdf")
        sys.exit(1)

    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        text = page.get_text().strip()
        if text:
            pages.append(text)
    doc.close()

    return "\n\n".join(pages), {"total_pages": len(pages), "format": "pdf"}


def extract_markdown(path: Path) -> Tuple[str, Dict]:
    """Extract text from Markdown file."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, {"format": "markdown"}


def extract_text(path: Path) -> Tuple[str, Dict]:
    """Extract text from plain text file."""
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text, {"format": "text"}


def extract_html(path: Path) -> Tuple[str, Dict]:
    """Extract text from HTML file."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        print("Install: pip install beautifulsoup4")
        sys.exit(1)

    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "html.parser")

    # Remove scripts, styles, nav, footer
    for tag in soup.find_all(["script", "style", "nav", "footer", "aside"]):
        tag.decompose()

    # Get title
    title = soup.title.string if soup.title else path.stem

    # Get main content
    main = soup.find("main") or soup.find("article") or soup.find("body")
    text = main.get_text(separator="\n", strip=True) if main else ""

    return text, {"format": "html", "title": title}


def extract_epub(path: Path) -> Tuple[str, Dict]:
    """Extract text from EPUB ebook."""
    try:
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup
    except ImportError:
        print("Install: pip install ebooklib beautifulsoup4")
        sys.exit(1)

    book = epub.read_epub(str(path))

    # Get metadata
    title = book.get_metadata("DC", "title")
    title = title[0][0] if title else path.stem

    author = book.get_metadata("DC", "creator")
    author = author[0][0] if author else None

    # Extract text from all documents
    chapters = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            soup = BeautifulSoup(item.get_content(), "html.parser")
            text = soup.get_text(separator="\n", strip=True)
            if text:
                chapters.append(text)

    return "\n\n".join(chapters), {
        "format": "epub",
        "title": title,
        "author": author,
        "chapters": len(chapters)
    }


def extract_docx(path: Path) -> Tuple[str, Dict]:
    """Extract text from Word document."""
    try:
        from docx import Document
    except ImportError:
        print("Install: pip install python-docx")
        sys.exit(1)

    doc = Document(str(path))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs), {"format": "docx"}


def extract_notebook(path: Path) -> Tuple[str, Dict]:
    """Extract text from Jupyter notebook."""
    content = json.loads(path.read_text(encoding="utf-8"))

    cells = []
    code_cells = 0
    markdown_cells = 0

    for cell in content.get("cells", []):
        cell_type = cell.get("cell_type", "")
        source = "".join(cell.get("source", []))

        if cell_type == "markdown":
            cells.append(source)
            markdown_cells += 1
        elif cell_type == "code":
            cells.append(f"```python\n{source}\n```")
            code_cells += 1

    return "\n\n".join(cells), {
        "format": "jupyter",
        "code_cells": code_cells,
        "markdown_cells": markdown_cells
    }


def extract_code(path: Path) -> Tuple[str, Dict]:
    """Extract text from code file."""
    text = path.read_text(encoding="utf-8", errors="ignore")

    # Detect language from extension
    ext_to_lang = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".go": "go",
        ".rs": "rust",
        ".java": "java",
        ".c": "c",
        ".cpp": "cpp",
        ".h": "c",
        ".hpp": "cpp",
        ".rb": "ruby",
        ".sh": "bash",
        ".bash": "bash",
        ".zsh": "zsh",
        ".sql": "sql",
        ".yaml": "yaml",
        ".yml": "yaml",
        ".json": "json",
        ".toml": "toml",
    }

    lang = ext_to_lang.get(path.suffix.lower(), "text")

    return text, {"format": "code", "language": lang}


# =============================================================================
# Format detection and routing
# =============================================================================

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".md": extract_markdown,
    ".markdown": extract_markdown,
    ".txt": extract_text,
    ".rst": extract_text,
    ".html": extract_html,
    ".htm": extract_html,
    ".epub": extract_epub,
    ".docx": extract_docx,
    ".ipynb": extract_notebook,
    # Code files
    ".py": extract_code,
    ".js": extract_code,
    ".ts": extract_code,
    ".go": extract_code,
    ".rs": extract_code,
    ".java": extract_code,
    ".c": extract_code,
    ".cpp": extract_code,
    ".h": extract_code,
    ".hpp": extract_code,
    ".rb": extract_code,
    ".sh": extract_code,
    ".bash": extract_code,
    ".sql": extract_code,
    ".yaml": extract_code,
    ".yml": extract_code,
}


def is_supported(path: Path) -> bool:
    """Check if file format is supported."""
    return path.suffix.lower() in EXTRACTORS


def extract(path: Path) -> Tuple[str, Dict]:
    """Extract text from file based on extension."""
    ext = path.suffix.lower()
    extractor = EXTRACTORS.get(ext)

    if not extractor:
        raise ValueError(f"Unsupported format: {ext}")

    return extractor(path)


# =============================================================================
# Chunking
# =============================================================================

def chunk_text(
    text: str,
    chunk_size: int = 400,
    overlap: int = 50
) -> List[Dict]:
    """
    Chunk text into overlapping segments.

    Args:
        text: Full text to chunk
        chunk_size: Target words per chunk
        overlap: Words to overlap between chunks

    Returns:
        List of chunk dicts with content and word_count
    """
    # Split into paragraphs
    paragraphs = [p.strip() for p in re.split(r'\n\s*\n', text) if p.strip()]

    chunks = []
    current = []
    current_words = 0

    for para in paragraphs:
        para_words = len(para.split())

        # If adding this paragraph exceeds limit, save current chunk
        if current_words + para_words > chunk_size and current:
            chunks.append({
                "content": "\n\n".join(current),
                "word_count": current_words
            })

            # Keep last paragraph for overlap
            if current and len(current[-1].split()) <= overlap:
                current = [current[-1]]
                current_words = len(current[-1].split())
            else:
                current = []
                current_words = 0

        current.append(para)
        current_words += para_words

    # Don't forget last chunk
    if current:
        chunks.append({
            "content": "\n\n".join(current),
            "word_count": current_words
        })

    return chunks


# =============================================================================
# Qdrant ingestion
# =============================================================================

def ingest_to_qdrant(
    chunks: List[Dict],
    file_path: Path,
    file_metadata: Dict,
    collection: str,
    qdrant_url: str = "http://localhost:6333"
) -> int:
    """
    Ingest chunks into Qdrant.

    Args:
        chunks: List of chunk dicts
        file_path: Original file path
        file_metadata: Metadata from extraction
        collection: Qdrant collection name
        qdrant_url: Qdrant server URL

    Returns:
        Number of chunks ingested
    """
    # Initialize clients
    client = QdrantClient(url=qdrant_url)
    embedder = TextEmbedding("sentence-transformers/all-MiniLM-L6-v2")

    # Vector name used by mcp-server-qdrant
    vector_name = "fast-all-minilm-l6-v2"

    # Ensure collection exists
    collections = [c.name for c in client.get_collections().collections]
    if collection not in collections:
        client.create_collection(
            collection_name=collection,
            vectors_config={
                vector_name: VectorParams(size=384, distance=Distance.COSINE)
            }
        )
        print(f"Created collection: {collection}")

    # Prepare points
    points = []
    contents = [c["content"] for c in chunks]

    print(f"Generating embeddings for {len(chunks)} chunks...")
    embeddings = list(embedder.embed(contents))

    # Generate file hash for deduplication
    file_hash = hashlib.md5(file_path.read_bytes()).hexdigest()[:12]

    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        # Create unique ID from file hash + chunk index
        point_id = hashlib.md5(f"{file_hash}_{i}".encode()).hexdigest()

        # Build metadata
        metadata = {
            "source": "local_file",
            "content_type": file_metadata.get("format", "text"),
            "file_path": str(file_path.absolute()),
            "filename": file_path.name,
            "harvested_at": datetime.now().isoformat(),
            "chunk_index": i,
            "total_chunks": len(chunks),
            "word_count": chunk["word_count"],
            **{k: v for k, v in file_metadata.items() if v is not None}
        }

        # Use named vector to match mcp-server-qdrant format
        points.append(PointStruct(
            id=point_id,
            vector={vector_name: embedding.tolist()},
            payload={
                "document": chunk["content"],
                "metadata": metadata
            }
        ))

    # Upsert in batches
    batch_size = 100
    for i in range(0, len(points), batch_size):
        batch = points[i:i + batch_size]
        client.upsert(collection_name=collection, points=batch)
        print(f"  Ingested {min(i + batch_size, len(points))}/{len(points)} chunks")

    return len(points)


# =============================================================================
# Main
# =============================================================================

def process_file(
    path: Path,
    collection: str,
    chunk_size: int,
    qdrant_url: str
) -> Dict:
    """Process a single file."""
    print(f"\nProcessing: {path.name}")

    # Extract text
    text, metadata = extract(path)

    if not text.strip():
        print(f"  Warning: No text extracted from {path.name}")
        return {"file": str(path), "status": "empty", "chunks": 0}

    # Chunk
    chunks = chunk_text(text, chunk_size=chunk_size)
    print(f"  Extracted {len(text.split())} words -> {len(chunks)} chunks")

    # Ingest
    ingested = ingest_to_qdrant(
        chunks=chunks,
        file_path=path,
        file_metadata=metadata,
        collection=collection,
        qdrant_url=qdrant_url
    )

    return {
        "file": str(path),
        "status": "success",
        "chunks": ingested,
        "format": metadata.get("format", "unknown")
    }


def main():
    parser = argparse.ArgumentParser(
        description="Ingest local files into Qdrant vector database"
    )
    parser.add_argument(
        "path",
        type=Path,
        help="File or directory to ingest"
    )
    parser.add_argument(
        "--collection",
        default=os.environ.get("COLLECTION_NAME", "personal_memories"),
        help="Qdrant collection name (default: personal_memories)"
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=400,
        help="Target words per chunk (default: 400)"
    )
    parser.add_argument(
        "--qdrant-url",
        default=os.environ.get("QDRANT_URL", "http://localhost:6333"),
        help="Qdrant server URL (default: http://localhost:6333)"
    )
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Recursively process directories"
    )

    args = parser.parse_args()

    # Collect files to process
    files = []
    if args.path.is_file():
        if is_supported(args.path):
            files.append(args.path)
        else:
            print(f"Unsupported format: {args.path.suffix}")
            sys.exit(1)
    elif args.path.is_dir():
        pattern = "**/*" if args.recursive else "*"
        for p in args.path.glob(pattern):
            if p.is_file() and is_supported(p):
                files.append(p)

        if not files:
            print(f"No supported files found in {args.path}")
            sys.exit(1)
    else:
        print(f"Path not found: {args.path}")
        sys.exit(1)

    print(f"Found {len(files)} file(s) to ingest")
    print(f"Collection: {args.collection}")
    print(f"Chunk size: {args.chunk_size} words")

    # Process files
    results = []
    for path in files:
        try:
            result = process_file(
                path=path,
                collection=args.collection,
                chunk_size=args.chunk_size,
                qdrant_url=args.qdrant_url
            )
            results.append(result)
        except Exception as e:
            print(f"  Error: {e}")
            results.append({
                "file": str(path),
                "status": "error",
                "error": str(e)
            })

    # Summary
    print("\n" + "=" * 50)
    print("Summary:")
    total_chunks = sum(r.get("chunks", 0) for r in results)
    success = sum(1 for r in results if r["status"] == "success")
    errors = sum(1 for r in results if r["status"] == "error")

    print(f"  Files processed: {len(results)}")
    print(f"  Successful: {success}")
    print(f"  Errors: {errors}")
    print(f"  Total chunks ingested: {total_chunks}")


if __name__ == "__main__":
    main()
